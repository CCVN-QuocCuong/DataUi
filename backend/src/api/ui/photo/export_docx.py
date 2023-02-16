from datetime import datetime
import json
import requests
import io
import os
from src.shared.common import S3_BUCKET_STORAGE_COC_FILES
from src.shared.common import successResponse, errorResponse 
from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm 
import re
import boto3  

file_path_docx = '/tmp/export_photo.docx'

def handler(event, context):
    """
    Definition:
        - Function to export two or six selected photos into a docx file.  
    
    Args:
        - event: Contains input list url (code) paramaters.
        - context: Default parameters of lambda function
    
    Returns:
      - errorResponse object if status code is equals 400 (Handling failed)
      - successResponse object if status code is equal 200 and url output docx file (Handling success)
    """
    try: 
        payload = json.loads(event['body'])
        
        print(f'payload: {payload}')
        
        __items = payload["items"]
        
        if __items is not None:
            if  len(__items) == 2:
                response_url = export_two_images(payload=payload)
            else:
                response_url = export_six_images(payload=payload) 
        else:
           return errorResponse(400, "Exception: not found data")
        
        return successResponse({"url": response_url})

    except Exception as e:
        return errorResponse(400, "Exception: {}".format(e))
    
    
def get_images(items): 
    """
    Definition:
        - Function to download attached image from FLCD App from API endpoint and code list 
    
    Args:
        - items: Contains input list url (code) paramaters. 
    
    Returns:
      - response: (dict) presigned url output to response from FLCD  
    """
    # for path in paths: 
    __url = f"{os.environ['FLCD_ENV_ENPOINT_DOWLOAD_FILES']}/media/download/list"
    
    print(f'URL: {__url}')
    
    # 2. download the data behind the __url
    response = requests.post(__url, json=items) 
    
    print(f'FLCD response data: {response}')
    
    if response.ok: 
        return (json.loads(response.content.decode("utf-8")))  
    else:
        return None
    
def export_two_images(payload):
    """
    Definition:
        - Function to export two selected photos into a docx file.  
    
    Args:
        - payload: Contains input paramaters:
            - title: (str) Title of document so user input on form
            - items: (list) photos code input for photo
    
    Returns:
      - response: (str) presigned url output to response from S3 
    """
    try:
        heading = payload['title']
        items = payload['items']
        result_files = get_images(items)
        print(f'result_files: {result_files}')
        if result_files is None:
            return None     
        
        images = [] 
        for url in result_files:
            response = requests.get(url, stream=True)
            image = io.BytesIO(response.content)
            images.append(image) 
       
        doc = Document() # create doc
        
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
    
        section.left_margin = Mm(20)
        section.right_margin = Mm(15)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(10)
        
        doc.add_heading(heading, 2) # add heading
       
        # iterate over each image in directory
        for i in range(len(images)): # show 2 rows of
            p = doc.add_paragraph()
            r = p.add_run()
            
            # add image so it creates a row with 4 images
            r.add_picture(images[i], width=Inches(6.5), height=Inches(3.71)) 
            r.add_text('                     ') # add space for image seperation
            r.add_text(remove_tags(items[i]['paragraph'])) 
                
        # save to file
        doc.save(file_path_docx)
        
        # call upload to S3 and return url
        file_name = 'export_photo_list_{0}.docx'.format(datetime.now().strftime('%Y%m%d%H%M')) 
        key='DOCX/{0}'.format(file_name)
                
        # push object in to Bucket S3
        response = upload_docx_s3(bucket=S3_BUCKET_STORAGE_COC_FILES, key=key) 
        
        return response
    
    except Exception as e:
        print("Execute function execute_insert_survey_data  error: {0}".format(e)) 
        
def export_six_images(payload):
    """
    Definition:
        - Function to export six selected photos into a docx file.  
    
    Args:
        - event: Contains input paramaters:
            - title: (str) Title of document so user input on form
            - items: (list) photos code input for photo
    
    Returns:
      - response: (str) presigned url output to response from S3 
    """
    heading = payload['title']
    items = payload['items']
    result_files = get_images(items)
    if result_files is None:
        return None     
    
    images = [] 
    for url in result_files:
        response = requests.get(url, stream=True)
        image = io.BytesIO(response.content)
        images.append(image) 
   
    # create doc
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    section.left_margin = Mm(20)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(10)
    
    # add heading
    doc.add_heading(heading, 2) 
    
    # create table with two rows and columns (Per row images are 3)
    table = doc.add_table(rows=0, cols=2, style="Table Grid")   
    for i in range(len(images)):   
        # add image to cell and align center
        if i % 2 == 0: 
            image_row = table.add_row() # add row to table for images
            cap_row = table.add_row() # add row to table for image name 
            paragraph = image_row.cells[0].paragraphs[0] 
            paragraph.add_run().add_picture(images[i], width=Inches(3.50), height=Inches(2.06))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # add caption to table
            cap_row.cells[0].text = remove_tags(items[i]['paragraph'])
            cap_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  
        else:
            # add image to cell and align center
            paragraph = image_row.cells[1].paragraphs[0] 
            paragraph.add_run().add_picture(images[i], width=Inches(3.50), height=Inches(2.06))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # add caption to table
            cap_row.cells[1].text = remove_tags(items[i]['paragraph'])
            cap_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  
        
    # save to file
    doc.save(file_path_docx)
    
    # call upload to S3 and return url
    file_name = 'export_photo_list_{0}.docx'.format(datetime.now().strftime('%Y%m%d%H%M')) 
    key='DOCX/{0}'.format(file_name)
            
    # push object in to Bucket S3
    response = upload_docx_s3(bucket=S3_BUCKET_STORAGE_COC_FILES, key=key) 
    
    return response

def upload_docx_s3(bucket, key):
    """
    Definition:
        - Function to upload file docx to S3 bucket 
    
    Args:
        - bucket: (str) bucket name of S3 need to save file.
        - key: Name of object to save in Bucket
    
    Returns:
      - response: (str) presigned url output to response from S3
    """
    try:
        
        client = boto3.client('s3') 
        
        # Write file sql into S3
        with open(file_path_docx, "rb") as f:  
            client.upload_fileobj(f, bucket, key)
            print('Upload file {0} success'.format(key)) 
         
        # get response url
        response = client.generate_presigned_url(
            ClientMethod='get_object',
            Params={
                    'Bucket': bucket,
                    'Key': key,
                    'ResponseContentDisposition': 'attachment'
                    },
                    ExpiresIn=3600
            )
        
        print(f'generate_presigned_url response: {response}')
            
        # return response URL output
        return response
    except Exception as e:
        print("Execute function upload_docx_s3 error: {0}".format(e)) 

def remove_tags(text):
    """
    Definition:
        - Function to remove special characters in string
    
    Args:
        - text: Contains input str paramater 
    
    Returns:
      - cleantext: (str) output after clean  
    """
    cleanr = re.compile('<.*?>')
    cleantext = re.sub(cleanr, '', text)
    return cleantext